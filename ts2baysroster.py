from docx import Document
import pandas as pd
import argparse

# Your CSV export must match the following column format or
# you can adjust which column matches which field
# Team names must be under a column named Team and that is not configurable
LAST='last'
FIRST='first'
GRADE='Grade Level'
BIRTH='birthdate'
CITY='city'
GENDER='gender'
# Adjust the following prefix to match a prefix that all your teams names start with
TEAM_PREFIX='Northboro'


lookup = {
    "team_name": {"row": 1, "col": 0},
    "team_grade": {"row": 1, "col": 6},
    "team_gender": {"row": 1, "col": 7},
    "last": {"row": 10, "col": 2},
    "first": {"row": 10, "col": 3},
    "grade": {"row": 10, "col": 6},
    "birth": {"row": 10, "col": 8},
    "town": {"row": 10, "col": 9}
}


def sanitize(name):
    illegal_characters = ['/', '?', '<', '>', '\\', ':', '*', '|', '(', ')']
    characters = list(name)
    for index, character in enumerate(characters):
        if characters[index] in illegal_characters:
            characters[index] = '~'
    name = ''.join(characters)
    return name


def create_roster(allteams_dataframe, team_dataframe, roster, team_name):
    main_table = roster.tables[0]
    #print_table(main_table)
    grade = ""
    gender = ""
    print(team_dataframe)
    for index, player in enumerate(team_dataframe.iterrows()):
        print("processing " + player[1][LAST] + " " + player[1][FIRST])
        main_table.cell(lookup['last']['row'] + index, lookup['last']['col']).text += player[1][LAST]
        main_table.cell(lookup['first']['row'] + index, lookup['first']['col']).text += player[1][FIRST]
        grade = player[1][GRADE]
        grade = grade.replace('Grade ', '')
        main_table.cell(lookup['grade']['row'] + index, lookup['grade']['col']).text += grade
        main_table.cell(lookup['birth']['row'] + index, lookup['birth']['col']).text += player[1][BIRTH].strftime('%Y')
        main_table.cell(lookup['town']['row'] + index, lookup['town']['col']).text += player[1][CITY]
        gender = get_gender(player[1][GENDER])
    main_table.cell(lookup['team_name']['row'], lookup['team_name']['col']).text += team_name
    main_table.cell(lookup['team_grade']['row'], lookup['team_grade']['col']).text += grade
    main_table.cell(lookup['team_gender']['row'], lookup['team_gender']['col']).text += gender
    roster_file_name = team_name + '_roster.docx'
    roster_file_name = sanitize(roster_file_name)
    print('Saving roster to : ' + roster_file_name)
    roster.save(roster_file_name)


def print_table(main_table):
    for row in main_table.rows:
        for index, cell in enumerate(row.cells):
            print(str(row._index) + ":" + str(index) + ':' + cell.text)


def get_gender(gender):
    if gender == 'Male':
        return 'Boys'
    if gender == 'Female':
        return 'Girls'
    print('No gender found returning NA')
    return 'NA'


def configure_parser():
    parser = argparse.ArgumentParser(description='Convert TeamSnap Team spreadsheet to BAYS roster doc')
    parser.add_argument('--teamExcelFile', type=str, nargs=1, required=True, help='Path to team excel file')
    return parser


if __name__ == "__main__":
    parser = configure_parser()
    args = parser.parse_args()
    print("using file : " + args.teamExcelFile[0])
    allteams_df = pd.read_excel(args.teamExcelFile[0])
    all_teams = allteams_df.Team.unique().tolist()
    #all_teams.remove('Team 1')
    bays_teams = [tname for tname in all_teams if isinstance(tname, str) and tname.startswith(TEAM_PREFIX)]
    print(all_teams)
    for team_name in bays_teams:
        print("processing team " + team_name)
        team_df = allteams_df.query("Team == @team_name")
        team_df_players_only = team_df.loc[team_df['role'] == 'Player']
        roster_template = Document("./roster_templates/bays_roster.docx")
        roster_template_25 = Document("./roster_templates/bays_roster_25.docx")
        if team_df_players_only.shape[0] > 20:
            roster_template = roster_template_25
        create_roster(allteams_df, team_df_players_only, roster_template, team_name)
    print("Roster generated")